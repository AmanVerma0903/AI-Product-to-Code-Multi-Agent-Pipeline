from fastapi import APIRouter, Depends, HTTPException, UploadFile, File
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select, delete
from app.db.session import get_db
from app.models.user import User
from app.models.project import Document
from app.models.vector_store import DocumentChunk
from app.core.security import get_current_user
from app.services.rag_service import RAGService
import aiofiles
import os
from pathlib import Path

router = APIRouter()

UPLOAD_DIR = Path("uploads")
UPLOAD_DIR.mkdir(exist_ok=True)
MAX_FILE_SIZE = 20 * 1024 * 1024  # 20MB

ALLOWED_EXTENSIONS = {".pdf", ".txt", ".docx"}

def allowed_file(filename: str) -> bool:
    """Check if file extension is allowed."""
    return Path(filename).suffix.lower() in ALLOWED_EXTENSIONS

@router.post("/{project_id}/upload")
async def upload_document(
    project_id: int,
    file: UploadFile = File(...),
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Upload a document to a project for RAG context."""
    
    # Validate file
    if not allowed_file(file.filename):
        raise HTTPException(
            status_code=400,
            detail=f"File type not allowed. Allowed: {', '.join(ALLOWED_EXTENSIONS)}"
        )
    
    if not file.size or file.size > MAX_FILE_SIZE:
        raise HTTPException(
            status_code=400,
            detail=f"File size must be between 1B and {MAX_FILE_SIZE / 1024 / 1024}MB"
        )
    
    try:
        # Read file content
        content = await file.read()
        
        # Extract text from different file types
        if file.filename.endswith('.txt'):
            text_content = content.decode('utf-8')
        elif file.filename.endswith('.pdf'):
            import PyPDF2
            from io import BytesIO
            pdf_reader = PyPDF2.PdfReader(BytesIO(content))
            text_content = "\n".join(page.extract_text() for page in pdf_reader.pages)
        elif file.filename.endswith('.docx'):
            from docx import Document as DocxDocument
            from io import BytesIO
            docx = DocxDocument(BytesIO(content))
            text_content = "\n".join(para.text for para in docx.paragraphs)
        else:
            raise HTTPException(status_code=400, detail="Unsupported file type")
        
        # Save file to disk
        file_path = UPLOAD_DIR / f"{project_id}_{file.filename}"
        async with aiofiles.open(file_path, 'wb') as f:
            await f.write(content)

        db_document = Document(
            project_id=project_id,
            filename=file.filename,
            file_path=str(file_path)
        )
        db.add(db_document)
        await db.flush()

        rag_service = RAGService(db)
        chunk_count = await rag_service.index_document_chunks(db_document.id, text_content)
        await db.refresh(db_document)

        return {
            "document_id": db_document.id,
            "filename": file.filename,
            "chunk_count": chunk_count,
            "status": "uploaded"
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Upload failed: {str(e)}")

@router.get("/{project_id}/search")
async def search_documents(
    project_id: int,
    query: str,
    limit: int = 5,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Search documents in a project using semantic similarity."""
    try:
        rag_service = RAGService(db)
        results = await rag_service.search(project_id, query, limit)
        
        return {
            "query": query,
            "results": results,
            "count": len(results)
        }
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Search failed: {str(e)}")

@router.delete("/{document_id}")
async def delete_document(
    document_id: int,
    current_user: User = Depends(get_current_user),
    db: AsyncSession = Depends(get_db)
):
    """Delete a document and its embeddings."""
    result = await db.execute(select(Document).filter(Document.id == document_id))
    document = result.scalar_one_or_none()
    
    if not document:
        raise HTTPException(status_code=404, detail="Document not found")
    
    try:
        # Delete file from disk
        if document.file_path and os.path.exists(document.file_path):
            os.remove(document.file_path)
        
        await db.execute(delete(DocumentChunk).where(DocumentChunk.document_id == document_id))
        await db.delete(document)
        await db.commit()
        
        return {"status": "deleted", "document_id": document_id}
        
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Delete failed: {str(e)}")
